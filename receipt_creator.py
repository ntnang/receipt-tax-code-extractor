import pandas as pd
from docx import Document
from vietnam_number import n2w
import locale
from datetime import datetime

def generate_receipts_from_template_with_filled_data():
    data_file = "data.xlsx"
    template_file = "template.docx"

    """Read replacements from Excel and return a dictionary."""
    df = pd.read_excel(data_file)  # Read the Excel file

    # Define starting index
    start_index = 1

    for index, row in df.iloc[start_index:].iterrows():
        row_dict = {i: value for i, value in enumerate(df.iloc[index])}
        replace_placeholders(template_file, f"receipt_{index}.docx", row_dict)

def replace_placeholders(doc_file, output_file, replacements):
    # sh.copy(doc_path, output_path)

    """Replace placeholders in a Word document."""
    doc = Document(doc_file)
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                placeholder = f"{{{key}}}"
                if placeholder in run.text:
                    paragraph.text = run.text.replace(placeholder, str(value))
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            placeholder = f"{{{key}}}"
                            if placeholder in run.text:
                                match key:
                                    case 4: run.text = run.text.replace(placeholder, format_date_vn(value))
                                    case 5: run.text = run.text.replace(placeholder, locale.format_string('%.0f đồng', abs(value), grouping=True))
                                    case _: run.text = run.text.replace(placeholder, str(value))
                            if "{#}" in run.text:
                                run.text = run.text.replace("{#}", n2w(str(replacements.get(5))) + " đồng")

    doc.save(output_file)
    print(f"File saved as {output_file}")

def format_date_vn(date_str):
    date_obj = datetime.strptime(date_str, "%d/%m/%Y")
    return f"ngày {date_obj.day} tháng {date_obj.month} năm {date_obj.year}"

# Set locale to Vietnamese (Vietnam)
locale.setlocale(locale.LC_ALL, 'vi_VN.UTF-8')

generate_receipts_from_template_with_filled_data()
